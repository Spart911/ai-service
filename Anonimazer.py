import shutil
import sys

# -*- coding: utf-8 -*-
from natasha import (
    Segmenter,
    MorphVocab,
    NewsEmbedding,
    NewsMorphTagger,
    NewsSyntaxParser,
    NewsNERTagger,
    PER,
    NamesExtractor,
    Doc, LOC
)
import re
from PIL import Image, ImageDraw
import easyocr
import os
from pdf2docx import Converter
from docx import Document

def string_to_list(input_string):
    words_list = input_string.split()
    return words_list
def list_to_string(input_list):
    result_string = ' '.join(map(str, input_list))
    return result_string
def replace_letters_on_match(list1, list2):
    result_list = []

    for item1 in list1:
        for item2 in list2:
            if item1 == item2:
                replaced_item = '*' * len(item1)
                result_list.append(replaced_item)
                break
        else:
            result_list.append(item1)

    return result_list
def read_docx(file_path):
    doc = Document(file_path)
    text_elements = [paragraph.text for paragraph in doc.paragraphs]
    return text_elements

def modify_and_save_docx(file_path, modified_elements):
    doc = Document()
    for element in modified_elements:
        doc.add_paragraph(element)

    doc.save(file_path)


def text_ai(List, result_list, ListExeption, ListADD):
    segmenter = Segmenter()
    morph_vocab = MorphVocab()
    emb = NewsEmbedding()

    morph_tagger = NewsMorphTagger(emb)
    syntax_parser = NewsSyntaxParser(emb)
    ner_tagger = NewsNERTagger(emb)
    names_extractor = NamesExtractor(morph_vocab)
    text = ", ".join(capitalize_first_letter_in_all(List))
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    doc.parse_syntax(syntax_parser)
    doc.tag_ner(ner_tagger)
    for span in doc.spans:
        if span.type == PER:
            result_list.append(span.text)
        if span.type == LOC:
            result_list.append(span.text)
    modified_list = find_missing_elements(replace_numbers_with_asterisks(List), List)
    result_list = ListExeption + filter_unique_elements(List, result_list) + modified_list
    result_list = remove_duplicates(result_list, ListExeption)
    result_list = ListADD + result_list
    return result_list
def filter_unique_elements(list1, list2):
    # Приводим элементы второго списка к нижнему регистру
    lower_list2 = [item.lower() for item in list2]

    # Фильтруем элементы первого списка, оставляя только те, которых нет во втором списке с учетом регистра
    result = [item for item in list1 if item.lower()  in lower_list2]

    return result
def find_missing_elements(list_with_asterisks, list_without_asterisks):
    result_list = []

    for item in list_without_asterisks:
        if '*' not in item and item not in list_with_asterisks:
            result_list.append(item)

    return result_list
def capitalize_first_letter_in_all(input_list):
    """
    Принимает список строк и возвращает список с первой буквой в верхнем регистре, а остальными в нижнем.
    """
    result_list = []

    for input_string in input_list:
        # Проверяем, все ли буквы в верхнем регистре
        if input_string.isupper():
            # Преобразуем первую букву в верхний регистр, а остальные в нижний
            result_list.append(input_string.capitalize())
        else:
            # Если не все буквы в верхнем регистре, добавляем строку без изменений
            result_list.append(input_string)

    return result_list
def is_valid_home_format(home_string):
    # Регулярное выражение для формата "д.178/1"
    date_pattern = re.compile(r'\b(?:д|кв)\.\d+(?:/\d+)?\b', flags=re.IGNORECASE)

    return bool(date_pattern.match(home_string))

def is_valid_date_format(date_string):
    date_pattern = re.compile(r'\b\d{2}.\d{2}.\d{4}(г\.)?\b|\b\d{2}.\d{2}.\d{4}')

    return bool(date_pattern.match(date_string))
def is_valid_cost_format(cost_string):
    cost_pattern = re.compile(r'\b\d+(?:р\.|р)\b')

    return bool(cost_pattern.match(cost_string))
def read_string_from_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            print(f"Строка успешно прочитана из файла: {file_path}")
            return content
    except Exception as e:
        print(f"Ошибка при чтении строки из файла: {e}")
        return None

def read_and_create_list(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Убираем скобки и разделяем текст на слова
    words = [word.strip('[],') for word in content.split()]

    # Удаляем повторяющиеся слова
    unique_words = list(set(words))

    return unique_words
def remove_duplicates(list1, list2):
    unique_list1 = list(set(list1) - set(list2))
    return unique_list1
def save_string_to_file(file_path, content):
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
        print(f"Строка успешно сохранена в файл: {file_path}")
    except Exception as e:
        print(f"Ошибка при сохранении строки в файл: {e}")
def replace_numbers_with_asterisks(my_list):
    result_list = []
    chek_ot = 0
    chek_front_ot_zak = 0
    chek_front_ot_fz = 0
    chek_cite_street = 0
    chek_state= 0
    for item in my_list:
        if chek_state == 0 and "." not in item  and "ФЗ"  not in item or (is_valid_date_format(item) and chek_ot == 0) or is_valid_home_format(item) or is_valid_cost_format(item):
            modified_item = ''.join('*' if c.isdigit() & c.isdigit() else c for c in item)
            if chek_cite_street == 1:
                modified_item = "*****"
            result_list.append(modified_item)
        else:
            if "@"  in item or "https://"  in item:
                modified_item = "*****"
                result_list.append(modified_item)
            else:
                result_list.append(item)
        if (item == "от" and chek_front_ot_zak == 1) or (item == "от" and chek_front_ot_fz == 1):
            chek_ot = 1
        else:
            chek_ot = 0

        if item == "ул." or item == "г." or item == "г.о." :
            chek_cite_street = 1
        else:
            chek_cite_street = 0



        if item ==  "ст." or item ==  "(ст.":
            chek_state = 1
        else:
            chek_state = 0



        if (item == "от" and chek_front_ot_zak == 1) or (item == "от" and chek_front_ot_fz == 1):
            chek_ot = 1
        else:
            chek_ot = 0
        if item == "закона":
            chek_front_ot_zak = 1
            chek_front_ot_fz = 0
        else:
            chek_front_ot_zak = 0

        if item == "ФЗ":
            chek_front_ot_fz = 1
            chek_front_ot_zak = 0
        else:
            chek_front_ot_fz = 0
    return result_list
def replace_elements_with_asterisks(list1, list2):
    for i in range(len(list1)):
        for element in list2:
            if element in list1[i]:
                list1[i] = list1[i].replace(element, '******')
def create_rectangles(image_path, result, words_to_hide, ListCoordinate):
    # Load the image
    img = Image.open(image_path)
    draw = ImageDraw.Draw(img)

    # Iterate through each word and check if it needs to be hidden
    for detection, coordinates in zip(result, ListCoordinate):
        word = detection
        # Check if the word needs to be hidden
        if any(hidden_word in word for hidden_word in words_to_hide):
            # Extract the correct coordinates format (x0, y0, x1, y1) from ListCoordinate
            x0, y0, x1, y1 = coordinates[0][0], coordinates[0][1], coordinates[2][0], coordinates[2][1]
            # Draw a rectangle to hide the word
            draw.rectangle([x0, y0, x1, y1], fill="black")

    # Save the image with hidden words outside the loop
    img.save(image_path)


def extract_text_from_image(image_path, List, ListCoordinate, languages=['en','ru']):
    # Initialize the EasyOCR Reader with the selected languages
    reader = easyocr.Reader(languages)

    # Get the result of text recognition
    result = reader.readtext(image_path)
    for detection in result:
        List.append(detection[1])
        ListCoordinate.append(detection[0])

def move_file(source_path, destination_path):
    try:
        shutil.move(source_path, destination_path)
        print(f'Файл успешно перемещен из {source_path} в {destination_path}')
        return destination_path
    except Exception as e:
        print(f'Ошибка при перемещении файла: {e}')
        return None
def rotate_image_90(image_path):
    # Открываем изображение
    img = Image.open(image_path)
    # Поворачиваем изображение на 90 градусов
    rotated_img = img.rotate(90)
    # Сохраняем результат
    rotated_img.save(image_path)
def dox_ai(my_list, result_list, ListExeption, ListADD):
    for i in range(len(my_list)):
        my_list[i] = string_to_list(my_list[i])
        result_abz = abz_ai(my_list[i], result_list, ListExeption, ListADD)
        my_list[i] = replace_letters_on_match(my_list[i], result_abz)
        my_list[i] = list_to_string(my_list[i])
    return my_list
def abz_ai(List, result_list, ListExeption, ListADD):
    segmenter = Segmenter()
    morph_vocab = MorphVocab()
    emb = NewsEmbedding()

    morph_tagger = NewsMorphTagger(emb)
    syntax_parser = NewsSyntaxParser(emb)
    ner_tagger = NewsNERTagger(emb)
    names_extractor = NamesExtractor(morph_vocab)
    text = ", ".join(capitalize_first_letter_in_all(List))
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    doc.parse_syntax(syntax_parser)
    doc.tag_ner(ner_tagger)
    for span in doc.spans:
        if span.type == PER:
            result_list.append(span.text)
        if span.type == LOC:
            result_list.append(span.text)
    modified_list = find_missing_elements(replace_numbers_with_asterisks(List), List)
    result_list = ListExeption + filter_unique_elements(List, result_list) + modified_list
    result_list = remove_duplicates(result_list, ListExeption)
    result_list = ListADD + result_list
    return result_list
# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    file_path_exeption = 'exeption.txt'
    directory_path = './../input'
    end_path = './../output/'
    List = []
    result_list = []
    ListADD = ["января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября",
               "ноября", "декабря", "мужской", "женский", "Мужской", "Женский"]
    ListExeption = read_and_create_list(file_path_exeption)

    files = os.listdir(directory_path)
    file_path = files[0]
    file_path = os.path.abspath(os.path.join(directory_path, file_path))
    print(f'Имя файла: {file_path}')
    if not os.path.isfile(file_path):
        print(f'Файл не найден: {file_path}')
    name_file = os.path.basename(file_path)
    # Проверяем разрешение по пути
    _, extension = os.path.splitext(file_path)

    if extension.lower() == '.pdf':
        # Если файл PDF, конвертируем его в DOCX
        docx_path = os.path.join(end_path, os.path.splitext(name_file)[0] + '.docx')
        cv = Converter(file_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        print(f'Конвертация успешно завершена: {file_path} -> {docx_path}')
        original_elements = read_docx(docx_path)
        print(original_elements)
        result_list = dox_ai(original_elements, result_list, ListExeption, ListADD)
        print(result_list)
        # Алгоритм модификации списка original_elements
        # Заменяем элементы, содержащие 'замена', на '*'
        # Сохраняем измененный список обратно в файл
        modify_and_save_docx(docx_path, result_list)

    elif extension.lower() in ['.png', '.jpg', '.jpeg']:
        ListCoordinate = []
        image_path = os.path.join(directory_path, name_file)
        end_image_path = os.path.join(end_path, name_file)
        extract_text_from_image(image_path, List, ListCoordinate)
        result_list = text_ai(List, result_list, ListExeption, ListADD)
        create_rectangles(image_path, List, result_list, ListCoordinate)
        image_path = move_file(image_path, end_image_path)
        # rotate_image_90(end_image_path)


    elif extension.lower() in ['.doc', '.docx']:
        docx_path = os.path.join(end_path, name_file)
        original_elements = read_docx(file_path)
        result_list = dox_ai(original_elements, result_list, ListExeption, ListADD)
        modify_and_save_docx(docx_path, result_list)
        os.remove(file_path)
    else:
        print(f'Неподдерживаемый формат файла: {file_path}')
    sys.exit()




